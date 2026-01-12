import streamlit as st
import google.generativeai as genai
import json
from datetime import date
import io
import requests
import urllib.parse
import random
import re

# --- NEW LIBRARY FOR WORD DOCS ---
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="DLP Generator", layout="centered")

# --- 2. API KEY EMBEDDED IN CODE ---
# Replace this with your actual Google AI API key
EMBEDDED_API_KEY = "AIzaSyCmtVgQOPR7htY6_ELzCuYEc_DWcLVkvYo"  # REPLACE WITH YOUR ACTUAL KEY

# --- 3. SIMPLIFIED HEADER WITHOUT LOGOS ---
def add_custom_header():
    """Add custom header with maroon background (NO LOGOS)"""
    
    st.markdown("""
    <style>
    .header-container {
        text-align: center;
        padding: 20px;
        margin-bottom: 25px;
        background-color: #800000; /* MAROON BACKGROUND */
        border-radius: 10px;
        box-shadow: 0 4px 12px rgba(128, 0, 0, 0.3);
        color: white;
    }
    .dept-name {
        font-size: 24px;
        font-weight: bold;
        color: white;
        margin: 0;
        text-shadow: 1px 1px 3px rgba(0,0,0,0.3);
    }
    .division-name {
        font-size: 20px;
        font-weight: bold;
        color: #FFD700; /* Gold color for contrast */
        margin: 8px 0;
    }
    .school-name {
        font-size: 28px;
        font-weight: bold;
        color: white;
        margin: 8px 0;
        text-transform: uppercase;
        letter-spacing: 1.5px;
    }
    .header-subtext {
        font-size: 15px;
        color: #FFD700; /* Gold color */
        margin-top: 8px;
        font-style: italic;
        font-weight: bold;
    }
    
    /* App title styling */
    .app-title {
        font-size: 32px;
        font-weight: bold;
        text-align: center;
        color: #800000; /* Maroon */
        margin: 15px 0 25px 0;
        padding: 10px;
        border-bottom: 3px solid #800000;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    </style>
    
    <div class="header-container">
        <p class="dept-name">DEPARTMENT OF EDUCATION REGION XI</p>
        <p class="division-name">DIVISION OF DAVAO DEL SUR</p>
        <p class="school-name">MANUAL NATIONAL HIGH SCHOOL</p>
        <p class="header-subtext">Kiblawan North District</p>
    </div>
    """, unsafe_allow_html=True)

# --- 4. AI GENERATOR ---
def clean_json_string(json_string):
    """Clean the JSON string by removing invalid characters and fixing common issues"""
    if not json_string:
        return json_string
    
    # Remove markdown code blocks
    json_string = re.sub(r'```json\s*', '', json_string)
    json_string = re.sub(r'```\s*', '', json_string)
    
    # Remove bullet points and other invalid characters
    json_string = json_string.replace('•', '-')  # Replace bullet points with hyphens
    json_string = json_string.replace('\u2022', '-')  # Unicode bullet
    json_string = json_string.replace('\u25cf', '-')  # Black circle bullet
    
    # Fix truncated strings (add closing quotes)
    json_string = re.sub(r':\s*$', '": ""', json_string)  # Fix truncated values at end of line
    
    # Fix unclosed quotes in the middle of JSON
    lines = json_string.split('\n')
    cleaned_lines = []
    
    for i, line in enumerate(lines):
        # Count quotes in the line
        quote_count = line.count('"')
        
        # If odd number of quotes, add a closing quote at the end
        if quote_count % 2 == 1 and ':' in line:
            # Find the last colon position
            last_colon_pos = line.rfind(':')
            if last_colon_pos > 0:
                # Check if there's an opening quote after the colon
                after_colon = line[last_colon_pos + 1:].strip()
                if after_colon.startswith('"') and not after_colon.endswith('"'):
                    line = line + '"'
                elif not after_colon.startswith('"') and after_colon:
                    # If value doesn't start with quote but should be string
                    value_start = last_colon_pos + 1
                    while value_start < len(line) and line[value_start] in ' \t':
                        value_start += 1
                    if value_start < len(line):
                        line = line[:value_start] + '"' + line[value_start:] + '"'
        
        cleaned_lines.append(line)
    
    json_string = '\n'.join(cleaned_lines)
    
    # Remove any control characters except newlines and tabs
    json_string = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', json_string)
    
    # Fix common JSON issues
    json_string = re.sub(r',\s*}', '}', json_string)  # Remove trailing commas before }
    json_string = re.sub(r',\s*]', ']', json_string)  # Remove trailing commas before ]
    
    return json_string

def generate_lesson_content(subject, grade, quarter, content_std, perf_std, competency, 
                           obj_cognitive=None, obj_psychomotor=None, obj_affective=None):
    try:
        # Use the embedded API key
        genai.configure(api_key=EMBEDDED_API_KEY)
        
        # Try multiple model options
        model_options = ['gemini-2.5-flash', 'gemini-1.5-flash', 'gemini-pro']
        model = None
        
        for model_name in model_options:
            try:
                model = genai.GenerativeModel(model_name)
                # Test with a simple prompt
                test_response = model.generate_content("Hello")
                if test_response:
                    st.sidebar.success(f"✓ Using model: {model_name}")
                    break
            except Exception as e:
                continue
        
        if not model:
            # Fallback to default
            model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Check if user provided objectives
        user_provided_objectives = obj_cognitive and obj_psychomotor and obj_affective
        
        if user_provided_objectives:
            # Use user-provided objectives in the prompt
            prompt = f"""
            You are an expert teacher from Manual National High School in the Division of Davao Del Sur, Region XI, Philippines.
            Create a JSON object for a Daily Lesson Plan (DLP).
            Subject: {subject}, Grade: {grade}, Quarter: {quarter}
            Content Standard: {content_std}
            Performance Standard: {perf_std}
            Learning Competency: {competency}

            USER-PROVIDED OBJECTIVES:
            - Cognitive: {obj_cognitive}
            - Psychomotor: {obj_psychomotor}
            - Affective: {obj_affective}

            IMPORTANT: Use these exact objectives provided by the user. Do NOT modify them.
            
            CRITICAL INSTRUCTIONS:
            1. You MUST generate exactly 5 distinct MULTIPLE CHOICE assessment questions with A, B, C, D choices.
            2. Each assessment question MUST follow this format: "question|A. choice1|B. choice2|C. choice3|D. choice4"
            3. The correct answer should be included in the choices.
            4. Return ONLY valid JSON format.
            5. Do NOT use bullet points (•) or any markdown in the JSON values.
            6. All string values must be properly quoted.
            7. Do NOT include any explanations outside the JSON.

            Return ONLY raw JSON. No markdown formatting.
            Structure:
            {{
                "obj_1": "Cognitive objective",
                "obj_2": "Psychomotor objective",
                "obj_3": "Affective objective",
                "topic": "The main topic (include math equations like 3x^2 if needed)",
                "integration_within": "Topic within same subject",
                "integration_across": "Topic across other subject",
                "resources": {{
                    "guide": "Teacher Guide reference",
                    "materials": "Learner Materials reference",
                    "textbook": "Textbook reference",
                    "portal": "Learning Resource Portal reference",
                    "other": "Other Learning Resources"
                }},
                "procedure": {{
                    "review": "Review activity",
                    "purpose_situation": "Real-life situation motivation description",
                    "visual_prompt": "A simple 3-word visual description. Example: 'Red Apple Fruit'. NO sentences.",
                    "vocabulary": "5 terms with definitions",
                    "activity_main": "Main activity description",
                    "explicitation": "Detailed explanation of the concept with clear explanations and TWO specific examples with detailed explanations",
                    "group_1": "Group 1 task",
                    "group_2": "Group 2 task",
                    "group_3": "Group 3 task",
                    "generalization": "Reflection questions"
                }},
                "evaluation": {{
                    "assess_q1": "Question 1 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q2": "Question 2 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q3": "Question 3 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q4": "Question 4 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q5": "Question 5 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assignment": "Assignment task",
                    "remarks": "Remarks",
                    "reflection": "Reflection"
                }}
            }}
            """
        else:
            # Generate objectives automatically
            prompt = f"""
            You are an expert teacher from Manual National High School in the Division of Davao Del Sur, Region XI, Philippines.
            Create a JSON object for a Daily Lesson Plan (DLP).
            Subject: {subject}, Grade: {grade}, Quarter: {quarter}
            Content Standard: {content_std}
            Performance Standard: {perf_std}
            Learning Competency: {competency}

            CRITICAL INSTRUCTIONS:
            1. You MUST generate exactly 5 distinct MULTIPLE CHOICE assessment questions with A, B, C, D choices.
            2. Each assessment question MUST follow this format: "question|A. choice1|B. choice2|C. choice3|D. choice4"
            3. The correct answer should be included in the choices.
            4. Return ONLY valid JSON format.
            5. Do NOT use bullet points (•) or any markdown in the JSON values.
            6. All string values must be properly quoted.
            7. Do NOT include any explanations outside the JSON.

            Return ONLY raw JSON. No markdown formatting.
            Structure:
            {{
                "obj_1": "Cognitive objective",
                "obj_2": "Psychomotor objective",
                "obj_3": "Affective objective",
                "topic": "The main topic (include math equations like 3x^2 if needed)",
                "integration_within": "Topic within same subject",
                "integration_across": "Topic across other subject",
                "resources": {{
                    "guide": "Teacher Guide reference",
                    "materials": "Learner Materials reference",
                    "textbook": "Textbook reference",
                    "portal": "Learning Resource Portal reference",
                    "other": "Other Learning Resources"
                }},
                "procedure": {{
                    "review": "Review activity",
                    "purpose_situation": "Real-life situation motivation description",
                    "visual_prompt": "A simple 3-word visual description. Example: 'Red Apple Fruit'. NO sentences.",
                    "vocabulary": "5 terms with definitions",
                    "activity_main": "Main activity description",
                    "explicitation": "Detailed explanation of the concept with clear explanations and TWO specific examples with detailed explanations",
                    "group_1": "Group 1 task",
                    "group_2": "Group 2 task",
                    "group_3": "Group 3 task",
                    "generalization": "Reflection questions"
                }},
                "evaluation": {{
                    "assess_q1": "Question 1 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q2": "Question 2 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q3": "Question 3 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q4": "Question 4 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assess_q5": "Question 5 with choices in format: question|A. choice1|B. choice2|C. choice3|D. choice4",
                    "assignment": "Assignment task",
                    "remarks": "Remarks",
                    "reflection": "Reflection"
                }}
            }}
            """
        
        response = model.generate_content(prompt)
        text = response.text
        
        # Clean the JSON response
        cleaned_text = clean_json_string(text)
        
        # Log for debugging
        st.sidebar.text_area("Raw AI Response", cleaned_text[:1000], height=200)
        
        # Try to parse the JSON
        try:
            ai_data = json.loads(cleaned_text)
            return ai_data
        except json.JSONDecodeError as je:
            st.error(f"JSON Parsing Error: {je}")
            st.sidebar.error("Failed to parse JSON. Attempting manual fix...")
            
            # Attempt manual extraction
            try:
                # Try to extract JSON using regex
                json_pattern = r'\{.*\}'
                match = re.search(json_pattern, cleaned_text, re.DOTALL)
                if match:
                    json_str = match.group(0)
                    # Remove any trailing commas
                    json_str = re.sub(r',\s*}', '}', json_str)
                    json_str = re.sub(r',\s*]', ']', json_str)
                    ai_data = json.loads(json_str)
                    return ai_data
            except Exception as e2:
                st.error(f"Manual JSON extraction also failed: {e2}")
                # Create fallback data
                return create_fallback_data(subject, grade, quarter, content_std, perf_std, competency)
        
    except Exception as e:
        st.error(f"AI Generation Error: {str(e)}")
        # Create fallback data
        return create_fallback_data(subject, grade, quarter, content_std, perf_std, competency)

def create_fallback_data(subject, grade, quarter, content_std, perf_std, competency):
    """Create fallback data in case AI generation fails"""
    return {
        "obj_1": f"Understand {subject} concepts",
        "obj_2": f"Apply {subject} skills",
        "obj_3": f"Appreciate the value of {subject}",
        "topic": f"Introduction to {subject}",
        "integration_within": f"Related {subject} topics",
        "integration_across": "Mathematics, Science",
        "resources": {
            "guide": "Teacher's Guide",
            "materials": "Learner's Materials",
            "textbook": f"{subject} Textbook",
            "portal": "DepEd LR Portal",
            "other": "Online resources"
        },
        "procedure": {
            "review": "Review previous lesson",
            "purpose_situation": "Real-world application",
            "visual_prompt": "Classroom Learning",
            "vocabulary": "Term1: Definition1\nTerm2: Definition2\nTerm3: Definition3\nTerm4: Definition4\nTerm5: Definition5",
            "activity_main": "Group activity to explore the topic",
            "explicitation": f"Detailed explanation of {subject} with examples. Example 1: Basic application. Example 2: Advanced application.",
            "group_1": "Research task",
            "group_2": "Problem-solving task",
            "group_3": "Presentation task",
            "generalization": "What did you learn? How can you apply this?"
        },
        "evaluation": {
            "assess_q1": f"What is the main concept of {subject}?|A. Concept A|B. Concept B|C. Concept C|D. Concept D",
            "assess_q2": f"How would you apply {subject} in real life?|A. Application A|B. Application B|C. Application C|D. Application D",
            "assess_q3": f"Explain the difference between key terms in {subject}.|A. Difference A|B. Difference B|C. Difference C|D. Difference D",
            "assess_q4": f"Solve a simple problem using {subject} concepts.|A. Solution A|B. Solution B|C. Solution C|D. Solution D",
            "assess_q5": f"What are the limitations of {subject} approaches?|A. Limitation A|B. Limitation B|C. Limitation C|D. Limitation D",
            "assignment": "Research more about the topic",
            "remarks": "Lesson delivered successfully",
            "reflection": "Students showed good understanding"
        }
    }

# --- 5. IMAGE FETCHER ---
def fetch_ai_image(keywords):
    if not keywords: keywords = "school_classroom"
    clean_prompt = re.sub(r'[\n\r\t]', ' ', str(keywords))
    clean_prompt = re.sub(r'[^a-zA-Z0-9 ]', '', clean_prompt).strip()
    
    encoded_prompt = urllib.parse.quote(clean_prompt)
    seed = random.randint(1, 9999)
    url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=600&height=350&nologo=true&seed={seed}"
    url = url.strip()
    
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except Exception:
        return None
    return None

# --- 6. DOCX HELPERS ---
def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_text(paragraph, text):
    """
    Parses text for ^ (superscript) and _ (subscript).
    """
    if not text:
        return

    pattern = r"([^\^_]*)(([\^_])([0-9a-zA-Z\-]+))(.*)"
    current_text = str(text)
    
    if "^" not in current_text and "_" not in current_text:
        paragraph.add_run(current_text)
        return

    while True:
        match = re.match(pattern, current_text)
        if match:
            pre_text = match.group(1)
            marker = match.group(3)
            script_text = match.group(4)
            rest = match.group(5)
            
            if pre_text:
                paragraph.add_run(pre_text)
            
            run = paragraph.add_run(script_text)
            if marker == '^':
                run.font.superscript = True
            elif marker == '_':
                run.font.subscript = True
                
            current_text = rest
            if not current_text:
                break
        else:
            paragraph.add_run(current_text)
            break

def add_row(table, label, content, bold_label=True):
    """Adds a row and applies formatting to the content."""
    row_cells = table.add_row().cells
    
    # Label Column (Left)
    p_lbl = row_cells[0].paragraphs[0]
    run_lbl = p_lbl.add_run(label)
    if bold_label:
        run_lbl.bold = True
    
    # Content Column (Right)
    text_content = ""
    
    if isinstance(content, list):
        # Join list items with newlines for vertical stacking
        text_content = "\n".join([str(item) for item in content])
    else:
        text_content = str(content) if content else ""
    
    format_text(row_cells[1].paragraphs[0], text_content)

def add_section_header(table, text):
    """Adds a full-width section header with Blue background."""
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_background(cell, "BDD7EE")

def parse_multiple_choice_question(q_text):
    """Parse a multiple choice question in format: question|A. choice1|B. choice2|C. choice3|D. choice4"""
    if not q_text:
        return "No question provided", []
    
    # Split by pipe character
    parts = q_text.split('|')
    
    if len(parts) < 5:
        # If not in expected format, return as-is
        return q_text, []
    
    question = parts[0].strip()
    choices = []
    
    # Extract choices (should be 4 choices)
    for i in range(1, min(5, len(parts))):
        choice = parts[i].strip()
        # Ensure choice starts with letter and period
        if not re.match(r'^[A-D]\.', choice):
            # Add prefix if missing
            choice_prefix = ['A.', 'B.', 'C.', 'D.'][i-1]
            choice = f"{choice_prefix} {choice}"
        choices.append(choice)
    
    # Ensure we have exactly 4 choices
    while len(choices) < 4:
        choices.append(f"{['A.', 'B.', 'C.', 'D.'][len(choices)]} Choice placeholder")
    
    return question, choices

def add_assessment_row(table, label, eval_sec):
    """Special function to add assessment row with multiple choice questions."""
    row_cells = table.add_row().cells
    
    # Label Column (Left)
    p_lbl = row_cells[0].paragraphs[0]
    run_lbl = p_lbl.add_run(label)
    run_lbl.bold = True
    
    # Content Column (Right) - Build from scratch
    content_cell = row_cells[1]
    
    # Clear cell by removing all existing paragraphs
    for paragraph in content_cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Create new content
    # 1. Header
    p_header = content_cell.add_paragraph()
    header_run = p_header.add_run("ASSESSMENT (5-item Multiple Choice Quiz)")
    header_run.bold = True
    
    # 2. Directions
    p_dir = content_cell.add_paragraph()
    p_dir.add_run("DIRECTIONS: Read each question carefully. Choose the letter of the correct answer from options A, B, C, and D.")
    
    # 3. Empty line for spacing
    content_cell.add_paragraph()
    
    # 4. Questions with multiple choice format
    for i in range(1, 6):
        question_key = f'assess_q{i}'
        raw_question = eval_sec.get(question_key, f'Question {i}')
        
        # Parse multiple choice question
        question_text, choices = parse_multiple_choice_question(raw_question)
        
        # Create question paragraph
        p_question = content_cell.add_paragraph()
        
        # Add question number (bold)
        num_run = p_question.add_run(f"{i}. ")
        num_run.bold = True
        
        # Add question text with formatting
        if question_text:
            format_text(p_question, question_text)
        
        # Add choices (A, B, C, D)
        if choices:
            for choice in choices:
                p_choice = content_cell.add_paragraph()
                p_choice.paragraph_format.left_indent = Inches(0.3)
                
                # Make the choice letter bold (A., B., etc.)
                choice_match = re.match(r'^([A-D]\.)\s*(.*)', choice)
                if choice_match:
                    letter_part = choice_match.group(1)
                    text_part = choice_match.group(2)
                    
                    letter_run = p_choice.add_run(f"{letter_part} ")
                    letter_run.bold = True
                    
                    if text_part:
                        format_text(p_choice, text_part)
                else:
                    # Fallback if format doesn't match
                    format_text(p_choice, choice)
        else:
            # Fallback: create placeholder choices
            for letter in ['A.', 'B.', 'C.', 'D.']:
                p_choice = content_cell.add_paragraph()
                p_choice.paragraph_format.left_indent = Inches(0.3)
                letter_run = p_choice.add_run(f"{letter} ")
                letter_run.bold = True
                p_choice.add_run(f"Choice {letter[0]}")
        
        # Add spacing between questions (except after last question)
        if i < 5:
            content_cell.add_paragraph()

# --- 7. DOCX CREATOR ---
def create_docx(inputs, ai_data, teacher_name, principal_name, uploaded_image):
    doc = Document()
    
    # --- SETUP A4 PAGE SIZE & MARGINS ---
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # --- FORMAL DOCUMENT HEADER ---
    # Main title: Republic of the Philippines
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Republic of the Philippines")
    title_run.font.size = Pt(24)
    title_run.font.name = "Times New Roman"
    title_run.bold = True
    
    # Subtitle: Department of Education
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Department of Education")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.name = "Arial"
    subtitle_run.bold = True
    
    # Body text lines
    body_lines = [
        "Region XI",
        "Schools Division Office of Davao del Sur",
        "Brey, Northern Paligue, Padada, Davao del Sur"
    ]
    
    for line in body_lines:
        body_para = doc.add_paragraph()
        body_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body_run = body_para.add_run(line)
        body_run.font.size = Pt(12)
        body_run.font.name = "Calibri"
    
    # Add spacing
    doc.add_paragraph()
    
    # Main document title
    doc_title_para = doc.add_paragraph()
    doc_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title_run = doc_title_para.add_run("Daily Lesson Log (DLL) / Daily Lesson Plan (DLP)")
    doc_title_run.font.size = Pt(14)
    doc_title_run.font.name = "Times New Roman"
    doc_title_run.bold = True
    
    # Add spacing
    doc.add_paragraph()

    # --- TOP INFO TABLE ---
    table_top = doc.add_table(rows=2, cols=4)
    table_top.style = 'Table Grid'
    table_top.autofit = False
    
    # Set column widths
    table_top.columns[0].width = Inches(2.0)
    table_top.columns[1].width = Inches(1.5)
    table_top.columns[2].width = Inches(1.5)
    table_top.columns[3].width = Inches(2.0)

    # Fill first row with labels
    row1 = table_top.rows[0].cells
    row1[0].text = "Subject Area:"
    row1[1].text = "Grade Level:"
    row1[2].text = "Quarter:"
    row1[3].text = "Date:"
    
    # Make labels bold
    for cell in row1:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Fill second row with values
    row2 = table_top.rows[1].cells
    format_text(row2[0].paragraphs[0], inputs['subject'])
    format_text(row2[1].paragraphs[0], inputs['grade'])
    format_text(row2[2].paragraphs[0], inputs['quarter'])
    format_text(row2[3].paragraphs[0], date.today().strftime('%B %d, %Y'))
    
    # Set font for values
    for cell in row2:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.name = "Calibri"
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Add spacing
    doc.add_paragraph()

    # --- MAIN CONTENT TABLE ---
    table_main = doc.add_table(rows=0, cols=2)
    table_main.style = 'Table Grid'
    table_main.autofit = False
    
    table_main.columns[0].width = Inches(2.0)
    table_main.columns[1].width = Inches(5.3)

    # Process Data
    objs = f"1. {ai_data.get('obj_1','')}\n2. {ai_data.get('obj_2','')}\n3. {ai_data.get('obj_3','')}"
    r = ai_data.get('resources', {})
    proc = ai_data.get('procedure', {})
    eval_sec = ai_data.get('evaluation', {})

    # SECTION I - with Roman numeral
    row_section1 = table_main.add_row().cells
    row_section1[0].merge(row_section1[1])
    section1_cell = row_section1[0]
    section1_cell.text = "I. CURRICULUM CONTENT, STANDARD AND LESSON COMPETENCIES"
    section1_cell.paragraphs[0].runs[0].bold = True
    section1_cell.paragraphs[0].runs[0].font.name = "Calibri"
    section1_cell.paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_background(section1_cell, "BDD7EE")
    
    # A. Content Standard
    add_row(table_main, "A. Content Standard", inputs['content_std'])
    
    # B. Performance Standard  
    add_row(table_main, "B. Performance Standard", inputs['perf_std'])
    
    # C. Learning Competencies and Objectives
    row_comp = table_main.add_row().cells
    row_comp[0].paragraphs[0].add_run("C. Learning Competencies and Objectives").bold = True
    row_comp[0].paragraphs[0].runs[0].font.name = "Calibri"
    row_comp[0].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Build the content for Learning Competencies
    p_comp = row_comp[1].paragraphs[0]
    p_comp.style.font.name = "Calibri"
    p_comp.style.font.size = Pt(11)
    
    # Learning Competency header
    comp_run = p_comp.add_run("Learning Competency\n")
    comp_run.bold = True
    comp_run.font.name = "Calibri"
    
    # Add the competency
    p_comp.add_run(f"{inputs['competency']}\n\n")
    
    # Learning Objectives header
    obj_run = p_comp.add_run("Learning Objectives:\n\n")
    obj_run.bold = True
    obj_run.font.name = "Calibri"
    
    # Add the objectives text
    p_comp.add_run("At the end of the lesson, the learners will be able to:\n\n")
    p_comp.add_run(objs)

    # D. Content
    add_row(table_main, "D. Content", ai_data.get('topic', ''))
    
    # E. Integration
    integration_text = f"Within: {ai_data.get('integration_within','')}\nAcross: {ai_data.get('integration_across','')}"
    add_row(table_main, "E. Integration", integration_text)

    # SECTION II - with Roman numeral
    row_section2 = table_main.add_row().cells
    row_section2[0].merge(row_section2[1])
    section2_cell = row_section2[0]
    section2_cell.text = "II. LEARNING RESOURCES"
    section2_cell.paragraphs[0].runs[0].bold = True
    section2_cell.paragraphs[0].runs[0].font.name = "Calibri"
    section2_cell.paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_background(section2_cell, "BDD7EE")
    
    # Learning Resources items
    add_row(table_main, "Teacher Guide", r.get('guide', ''))
    add_row(table_main, "Learner's Materials(LMs)", r.get('materials', ''))
    add_row(table_main, "Textbooks", r.get('textbook', ''))
    add_row(table_main, "Learning Resource (LR) Portal", r.get('portal', ''))
    add_row(table_main, "Other Learning Resources", r.get('other', ''))

    # SECTION III - with Roman numeral
    row_section3 = table_main.add_row().cells
    row_section3[0].merge(row_section3[1])
    section3_cell = row_section3[0]
    section3_cell.text = "III. TEACHING AND LEARNING PROCEDURE"
    section3_cell.paragraphs[0].runs[0].bold = True
    section3_cell.paragraphs[0].runs[0].font.name = "Calibri"
    section3_cell.paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_background(section3_cell, "BDD7EE")
    
    # A. Activating Prior Knowledge - with subtitle in parentheses
    review_content = f"{proc.get('review', '')}\n\n(Minds and Moods)"
    add_row(table_main, "A. Activating Prior Knowledge", review_content)
    
    # B. Establishing Lesson Purpose - with subtitle and image
    row_purpose = table_main.add_row().cells
    row_purpose[0].paragraphs[0].add_run("B. Establishing Lesson Purpose").bold = True
    row_purpose[0].paragraphs[0].runs[0].font.name = "Calibri"
    row_purpose[0].paragraphs[0].runs[0].font.size = Pt(11)
    
    purpose_cell = row_purpose[1]
    
    # Add subtitle
    purpose_p1 = purpose_cell.add_paragraph()
    purpose_p1.add_run("1. Lesson Purpose\n").bold = True
    purpose_p1.runs[0].font.name = "Calibri"
    
    # Add purpose situation
    if proc.get('purpose_situation', ''):
        purpose_para = purpose_cell.add_paragraph(proc.get('purpose_situation', ''))
        purpose_para.style.font.name = "Calibri"
        purpose_para.style.font.size = Pt(11)
    
    purpose_cell.add_paragraph()  # Add spacing
    
    # Add image
    img_data = None
    if uploaded_image:
        img_data = uploaded_image
    else:
        raw_prompt = proc.get('visual_prompt', 'school')
        img_data = fetch_ai_image(raw_prompt)
    
    if img_data:
        try:
            p_i = purpose_cell.add_paragraph()
            p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_i = p_i.add_run()
            run_i.add_picture(img_data, width=Inches(3.5))
        except:
            error_para = purpose_cell.add_paragraph("[Image Error]")
            error_para.style.font.name = "Calibri"
    else:
        noimg_para = purpose_cell.add_paragraph("[No Image Available]")
        noimg_para.style.font.name = "Calibri"
    
    purpose_cell.add_paragraph()  # Add spacing
    
    # Add vocabulary
    purpose_p2 = purpose_cell.add_paragraph()
    purpose_p2.add_run("2. Unlocking Content Vocabulary\n").bold = True
    purpose_p2.runs[0].font.name = "Calibri"
    
    if proc.get('vocabulary', ''):
        vocab_para = purpose_cell.add_paragraph(proc.get('vocabulary', ''))
        vocab_para.style.font.name = "Calibri"
        vocab_para.style.font.size = Pt(11)
    
    # Add subtitle at the end
    aims_para = purpose_cell.add_paragraph("\n(Aims)")
    aims_para.style.font.name = "Calibri"

    # C. Developing and Deepening Understanding - with subtitle
    developing_content = f"Activity: {proc.get('activity_main','')}\n\n"
    developing_content += f"Explicitation: {proc.get('explicitation','')}\n\n"
    developing_content += f"Working by group:\n\n"
    developing_content += f"Group 1: {proc.get('group_1','')}\n\n"
    developing_content += f"Group 2: {proc.get('group_2','')}\n\n"
    developing_content += f"Group 3: {proc.get('group_3','')}\n\n"
    developing_content += "(Tasks and Thought)"
    
    add_row(table_main, "C. Developing and Deepening Understanding", developing_content)
    
    # D. Making Generalization - with subtitle
    generalization_content = f"{proc.get('generalization', '')}\n\n(Abstract)"
    add_row(table_main, "D. Making Generalization", generalization_content)

    # SECTION IV - with Roman numeral
    row_section4 = table_main.add_row().cells
    row_section4[0].merge(row_section4[1])
    section4_cell = row_section4[0]
    section4_cell.text = "IV. EVALUATING LEARNING"
    section4_cell.paragraphs[0].runs[0].bold = True
    section4_cell.paragraphs[0].runs[0].font.name = "Calibri"
    section4_cell.paragraphs[0].runs[0].font.size = Pt(11)
    set_cell_background(section4_cell, "BDD7EE")
    
    # A. Tests/Assessment - with multiple choice questions
    add_assessment_row(table_main, "A. Tests/Assessment", eval_sec)
    
    # B. Assignment
    add_row(table_main, "B. Assignment", eval_sec.get('assignment', ''))
    
    # C. Teacher's Remarks
    add_row(table_main, "C. Teacher's Remarks", eval_sec.get('remarks', ''))
    
    # D. Reflection - with subtitle
    reflection_content = f"{eval_sec.get('reflection', '')}\n\n(Gains)"
    add_row(table_main, "D. Reflection", reflection_content)

    # Add spacing before signatures
    doc.add_paragraph()
    doc.add_paragraph()

    # --- SIGNATORIES TABLE ---
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.autofit = False
    
    sig_table.columns[0].width = Inches(3.0)
    sig_table.columns[1].width = Inches(3.0)
    sig_table.columns[2].width = Inches(3.0)
    
    row = sig_table.rows[0]
    
    # LEFT COLUMN: PREPARED BY
    prepared_cell = row.cells[0]
    prepared_p = prepared_cell.add_paragraph()
    prepared_p.add_run("Prepared by:\n\n").bold = True
    prepared_p.runs[0].font.name = "Calibri"
    
    # Add name
    name_p = prepared_cell.add_paragraph()
    name_run = name_p.add_run(teacher_name)
    name_run.bold = True
    name_run.font.name = "Calibri"
    
    # Add teacher position
    position_p = prepared_cell.add_paragraph("\nTeacher III")
    position_p.runs[0].font.name = "Calibri"
    
    # MIDDLE COLUMN: CHECKED BY (1)
    checked1_cell = row.cells[1]
    checked1_p = checked1_cell.add_paragraph()
    checked1_p.add_run("Checked by:\n\n").bold = True
    checked1_p.runs[0].font.name = "Calibri"
    
    # Add underline for first checker
    name1_p = checked1_cell.add_paragraph()
    name1_run = name1_p.add_run("KRYZLYN A. AMAR")
    name1_run.bold = True
    name1_run.font.name = "Calibri"
    
    # Add position
    position1_p = checked1_cell.add_paragraph("\nMaster Teacher I")
    position1_p.runs[0].font.name = "Calibri"
    
    # RIGHT COLUMN: NOTED BY
    noted_cell = row.cells[2]
    noted_p = noted_cell.add_paragraph()
    noted_p.add_run("Noted by:\n\n").bold = True
    noted_p.runs[0].font.name = "Calibri"
    
    # Add name for principal
    name2_p = noted_cell.add_paragraph()
    name2_run = name2_p.add_run(principal_name)
    name2_run.bold = True
    name2_run.font.name = "Calibri"
    
    # Add position
    position2_p = noted_cell.add_paragraph("\nSchool Head")
    position2_p.runs[0].font.name = "Calibri"

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 8. STREAMLIT UI ---
def main():
    # Add custom header with maroon background (NO LOGOS)
    add_custom_header()
    
    # App Title - IN ONE LINE with custom styling
    st.markdown('<p class="app-title">Daily Lesson Plan (DLP) Generator</p>', unsafe_allow_html=True)
    
    with st.sidebar:
        st.header("📋 User Information")
        
        # Set default names to the required values
        teacher_name = st.text_input("Teacher Name", value="RICHARD P. SAMORANOS")
        principal_name = st.text_input("Principal Name", value="ROSALITA A. ESTROPIA")
        
        st.markdown("---")
        st.info("Upload an image (optional) for the lesson")
        uploaded_image = st.file_uploader("Choose an image for lesson", type=['png', 'jpg', 'jpeg'], key="lesson")
        
        st.markdown("---")
        st.success("✅ API Key is already configured in the application")
    
    # Form Inputs
    col1, col2, col3 = st.columns(3)
    with col1:
        subject = st.text_input("Subject Area", placeholder="e.g., Mathematics")
    
    with col2:
        # Grade Level Dropdown - Kinder to Grade 12
        grade_options = [
            "Kinder",
            "Grade 1", "Grade 2", "Grade 3", "Grade 4", "Grade 5", "Grade 6",
            "Grade 7", "Grade 8", "Grade 9", "Grade 10",
            "Grade 11", "Grade 12"
        ]
        grade = st.selectbox("Grade Level", grade_options, index=9)  # Default to Grade 10
    
    with col3:
        # Quarter Dropdown - Roman Numerals
        quarter_options = ["I", "II", "III", "IV"]
        quarter = st.selectbox("Quarter", quarter_options, index=0)  # Default to Quarter I
    
    content_std = st.text_area("Content Standard", placeholder="The learner demonstrates understanding of...")
    perf_std = st.text_area("Performance Standard", placeholder="The learner is able to...")
    competency = st.text_area("Learning Competency", placeholder="Competency code and description...")
    
    st.markdown("---")
    
    # --- OPTIONAL LESSON OBJECTIVES SECTION ---
    st.subheader("📝 Optional: Lesson Objectives")
    st.info("If you already have your lesson objectives, enter them below. Otherwise, leave blank and AI will generate them.")
    
    with st.expander("Enter Lesson Objectives (Optional)", expanded=False):
        col_obj1, col_obj2, col_obj3 = st.columns(3)
        
        with col_obj1:
            obj_cognitive = st.text_area(
                "Cognitive Objective",
                placeholder="e.g., Identify the parts of a cell",
                height=100,
                help="What students should know or understand"
            )
        
        with col_obj2:
            obj_psychomotor = st.text_area(
                "Psychomotor Objective",
                placeholder="e.g., Draw and label the parts of a cell",
                height=100,
                help="What students should be able to do"
            )
        
        with col_obj3:
            obj_affective = st.text_area(
                "Affective Objective",
                placeholder="e.g., Appreciate the complexity of living organisms",
                height=100,
                help="Values, attitudes, or emotions to develop"
            )
    
    st.markdown("---")
    
    # Generate Button
    if st.button("🚀 Generate DLP", type="primary", use_container_width=True):
        if not all([subject, grade, quarter, content_std, perf_std, competency]):
            st.error("Please fill all required fields")
            return
        
        # Check if user provided objectives
        user_provided_objectives = obj_cognitive and obj_psychomotor and obj_affective
        
        if user_provided_objectives:
            st.info("✅ Using your provided lesson objectives")
            with st.spinner("🤖 Generating lesson content with YOUR objectives..."):
                ai_data = generate_lesson_content(
                    subject, grade, quarter, 
                    content_std, perf_std, competency,
                    obj_cognitive, obj_psychomotor, obj_affective
                )
        else:
            st.info("🔧 AI will generate lesson objectives for you")
            with st.spinner("🤖 Generating complete lesson content with AI..."):
                ai_data = generate_lesson_content(
                    subject, grade, quarter, 
                    content_std, perf_std, competency
                )
            
        if ai_data:
            st.success("✅ AI content generated successfully!")
            
            # Show objectives preview
            st.subheader("📋 Generated Objectives")
            col_obj_pre1, col_obj_pre2, col_obj_pre3 = st.columns(3)
            
            with col_obj_pre1:
                st.info("**Cognitive**")
                st.write(ai_data.get('obj_1', 'N/A'))
            
            with col_obj_pre2:
                st.info("**Psychomotor**")
                st.write(ai_data.get('obj_2', 'N/A'))
            
            with col_obj_pre3:
                st.info("**Affective**")
                st.write(ai_data.get('obj_3', 'N/A'))
            
            # Show assessment preview
            with st.expander("📝 Preview Assessment Questions"):
                for i in range(1, 6):
                    question_key = f'assess_q{i}'
                    raw_question = ai_data.get('evaluation', {}).get(question_key, '')
                    if raw_question:
                        question_text, choices = parse_multiple_choice_question(raw_question)
                        st.markdown(f"**Question {i}:** {question_text}")
                        if choices:
                            for choice in choices:
                                st.write(f"  {choice}")
                        st.markdown("---")
            
            # Full preview
            with st.expander("📄 Preview All Generated Content"):
                st.json(ai_data)
            
            # Create DOCX
            inputs = {
                'subject': subject,
                'grade': grade,
                'quarter': quarter,
                'content_std': content_std,
                'perf_std': perf_std,
                'competency': competency
            }
            
            with st.spinner("📄 Creating DOCX file..."):
                docx_buffer = create_docx(inputs, ai_data, teacher_name, principal_name, uploaded_image)
            
            # Download button
            st.download_button(
                label="📥 Download DLP (.docx)",
                data=docx_buffer,
                file_name=f"DLP_{subject}_{grade}_Q{quarter}_{date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            # Display success message
            st.balloons()
            st.success(f"✅ DLP generated for {subject} - {grade} - Quarter {quarter}")
        else:
            st.error("Failed to generate AI content. Please try again.")

if __name__ == "__main__":
    main()
